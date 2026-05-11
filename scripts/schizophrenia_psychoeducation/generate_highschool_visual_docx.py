#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
統合失調症 高校生向け図解版 .docx 生成スクリプト

matplotlib でチャート・図を生成し python-docx に埋め込む。

使用例:
  python3 generate_highschool_visual_docx.py
  python3 generate_highschool_visual_docx.py --output samples/output/my_output.docx
"""

from __future__ import annotations

import argparse
import io
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

# ---------------------------------------------------------------------------
# フォント・サイズ定数
# ---------------------------------------------------------------------------
BODY_FONT = "MS ゴシック"
PT_TITLE = Pt(16)
PT_H1 = Pt(13)
PT_H2 = Pt(11.5)
PT_BODY = Pt(10.5)
PT_SMALL = Pt(9.5)
PT_CAPTION = Pt(9)
LINE_SPACING = 1.35

# ---------------------------------------------------------------------------
# カラー定数 (hex, RGB)
# ---------------------------------------------------------------------------
C_PRIMARY = "1E3A5F"
C_TEAL = "0F766E"
C_HEADER_BG = "E8EEF5"
C_RED_LIGHT = "FEE2E2"
C_RED_MID = "FECACA"
C_BLUE_LIGHT = "DBEAFE"
C_BLUE_MID = "BFDBFE"
C_PURPLE_LIGHT = "EDE9FE"
C_PURPLE_MID = "DDD6FE"
C_GREEN_LIGHT = "DCFCE7"
C_AMBER_LIGHT = "FEF3C7"
C_ROW_ALT = "F3F4F7"
C_MYTH_LEFT = "FFF1F2"
C_MYTH_RIGHT = "F0FDF4"
C_BORDER_LIGHT = "E2E8F0"
C_STAT1_BG = "EFF6FF"
C_STAT2_BG = "F0FDFA"
C_STAT3_BG = "EEF2FF"

# ---------------------------------------------------------------------------
# ユーティリティ関数
# ---------------------------------------------------------------------------

def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), name)


def body_rhythm(para, *, tight: bool = False) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.widow_control = True
    pf.space_after = Pt(3 if tight else 5)


def add_run(para, text: str, *, bold: bool = False, size: Pt = PT_BODY,
            color_hex: str | None = None, italic: bool = False) -> None:
    r = para.add_run(text)
    set_east_asia_font(r, BODY_FONT)
    r.font.size = size
    r.font.bold = bold
    r.font.italic = italic
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)


def add_para(doc: Document, text: str = "", *, bold: bool = False,
             size: Pt = PT_BODY, tight: bool = False,
             align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
             color_hex: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = align
    body_rhythm(p, tight=tight)
    add_run(p, text, bold=bold, size=size, color_hex=color_hex)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, BODY_FONT)
    r.font.bold = True
    r.font.size = PT_H1
    r.font.color.rgb = RGBColor.from_string(C_PRIMARY)
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    pf.widow_control = True
    # 下線ボーダー
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:space"), "3")
    bot.set(qn("w:color"), C_TEAL)
    pbdr.append(bot)
    ppr.append(pbdr)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, BODY_FONT)
    r.font.bold = True
    r.font.size = PT_H2
    r.font.color.rgb = RGBColor.from_string(C_TEAL)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(5)
    pf.keep_with_next = True
    pf.widow_control = True


def set_cell_shading(cell, fill_hex: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcpr.append(shd)


def set_cell_padding(cell, dxa: int = 100) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    old = tcpr.find(qn("w:tcMar"))
    if old is not None:
        tcpr.remove(old)
    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        n = OxmlElement(f"w:{side}")
        n.set(qn("w:w"), str(dxa))
        n.set(qn("w:type"), "dxa")
        mar.append(n)
    tcpr.append(mar)


def set_cell_border(cell, color_hex: str = "E2E8F0", sz: str = "6") -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tcpr.append(borders)


def cell_para(cell, text: str, *, bold: bool = False, size: Pt = PT_BODY,
              color_hex: str | None = None, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
              italic: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.widow_control = True
    add_run(p, text, bold=bold, size=size, color_hex=color_hex, italic=italic)


def insert_image(doc: Document, buf: io.BytesIO, width_cm: float = 14.0) -> None:
    """BytesIO の PNG 画像を doc に挿入。"""
    buf.seek(0)
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run()
    r.add_picture(buf, width=Cm(width_cm))


def add_colored_box(doc: Document, title: str, body: str,
                    bg_hex: str, border_hex: str, title_hex: str) -> None:
    """1行タイトル + 本文のカラーボックス（1列1行テーブルで疑似実装）。"""
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # タイトル行
    th = table.rows[0].cells[0]
    set_cell_shading(th, border_hex)
    set_cell_padding(th, 90)
    cell_para(th, title, bold=True, size=PT_H2, color_hex="FFFFFF",
              align=WD_ALIGN_PARAGRAPH.LEFT)
    # 本文行
    tb = table.rows[1].cells[0]
    set_cell_shading(tb, bg_hex)
    set_cell_padding(tb, 110)
    cell_para(tb, body, size=PT_BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# matplotlib 図生成
# ---------------------------------------------------------------------------

def fig_age_bar() -> io.BytesIO:
    """図1: 年代別発症頻度 横棒グラフ。"""
    ages = ["55歳〜", "45–54歳", "35–44歳", "25–34歳", "15–24歳", "〜14歳"]
    values = [1.2, 2.0, 3.5, 5.0, 4.0, 1.0]
    colors = ["#93C5FD", "#60A5FA", "#3B82F6", "#1D4ED8", "#1E40AF", "#BFDBFE"]
    colors_rev = list(reversed(colors))

    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    bars = ax.barh(ages, values, color=colors_rev, edgecolor="white", linewidth=0.8,
                   height=0.6)
    ax.set_xlabel("相対的な発症頻度（概略）", fontsize=9, color="#475569")
    ax.set_xlim(0, 6.2)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.spines["bottom"].set_color("#CBD5E1")
    ax.tick_params(colors="#334155", labelsize=9)
    ax.set_facecolor("#F8FAFC")
    fig.patch.set_facecolor("#F8FAFC")

    for bar, val in zip(bars, values):
        ax.text(val + 0.1, bar.get_y() + bar.get_height() / 2,
                f"{val}", va="center", ha="left", fontsize=8.5, color="#1E3A5F",
                fontweight="bold")

    # ピーク強調
    ax.annotate("← ピーク（10代後半〜30代）",
                xy=(5.0, 2), xytext=(3.8, 1.0),
                arrowprops=dict(arrowstyle="->", color="#EF4444", lw=1.2),
                fontsize=8, color="#EF4444", fontweight="bold")

    ax.set_title("発症しやすい年代（相対的なイメージ）", fontsize=10,
                 color="#1E3A5F", fontweight="bold", pad=8)
    plt.tight_layout(pad=0.8)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    return buf


def fig_recovery_curve() -> io.BytesIO:
    """図2: 回復の波ライン グラフ。"""
    x = np.linspace(0, 10, 500)

    def recovery(t: np.ndarray) -> np.ndarray:
        base = 0.45 * t / 10
        wave = 0.35 * np.exp(-0.3 * t) * np.sin(2.5 * t + 1.0)
        return base + wave + 0.08

    y = recovery(x)

    fig, ax = plt.subplots(figsize=(7.5, 3.2))

    # フェーズ背景
    ax.axvspan(0, 2.5, alpha=0.12, color="#EF4444", label="急性期")
    ax.axvspan(2.5, 5.5, alpha=0.10, color="#F59E0B", label="回復期")
    ax.axvspan(5.5, 8.0, alpha=0.10, color="#3B82F6", label="安定期")
    ax.axvspan(8.0, 10, alpha=0.10, color="#10B981", label="生活の再建")

    # 曲線
    ax.plot(x, y, color="#0F766E", linewidth=2.2, zorder=5)
    ax.fill_between(x, 0, y, alpha=0.12, color="#0F766E")

    # フェーズラベル
    for xp, label, col in [
        (1.25, "急性期", "#EF4444"),
        (4.0,  "回復期", "#D97706"),
        (6.75, "安定期", "#2563EB"),
        (9.0,  "生活の\n再建", "#059669"),
    ]:
        ax.text(xp, -0.12, label, ha="center", va="top", fontsize=8.5,
                color=col, fontweight="bold")

    ax.set_ylabel("生活の質・機能レベル", fontsize=9, color="#475569")
    ax.set_xlim(0, 10)
    ax.set_ylim(-0.05, 0.95)
    ax.set_yticks([0.0, 0.5, 0.9])
    ax.set_yticklabels(["低", "中", "高"], fontsize=8.5, color="#64748B")
    ax.set_xticks([])
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["bottom"].set_visible(False)
    ax.spines["left"].set_color("#CBD5E1")
    ax.set_facecolor("#F8FAFC")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_title("回復のイメージ（波を繰り返しながら改善していく）", fontsize=10,
                 color="#1E3A5F", fontweight="bold", pad=8)

    plt.tight_layout(pad=0.8)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    return buf


def fig_cause_radial() -> io.BytesIO:
    """図3: 原因要因 放射図。"""
    fig, ax = plt.subplots(figsize=(7.5, 4.0))
    ax.set_xlim(-1.1, 1.1)
    ax.set_ylim(-1.1, 1.1)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")

    # 中央円
    center = mpatches.Circle((0, 0), 0.22, color="#1E3A5F", zorder=5)
    ax.add_patch(center)
    ax.text(0, 0, "統合\n失調症", ha="center", va="center", fontsize=10,
            color="white", fontweight="bold", zorder=6)

    # 周囲バブル
    factors = [
        ( 0.0,  0.75, "脳の神経伝達\n物質のバランス", "#DBEAFE", "#3B82F6", "#1D4ED8"),
        ( 0.75,  0.0, "遺  伝\n的要因",            "#DCFCE7", "#22C55E", "#15803D"),
        ( 0.0, -0.75, "強いストレス\n・睡眠不足",   "#FEF3C7", "#F59E0B", "#B45309"),
        (-0.75,  0.0, "環境・脳の\n発  達",         "#F3E8FF", "#A855F7", "#7C3AED"),
    ]

    for x, y, label, bg, edge, txt in factors:
        # 接続線
        ax.annotate("", xy=(x * 0.24, y * 0.24),
                    xytext=(x * 0.52, y * 0.52),
                    arrowprops=dict(arrowstyle="-", color="#94A3B8", lw=1.5))
        # バブル
        bubble = mpatches.FancyBboxPatch(
            (x - 0.225, y - 0.14), 0.45, 0.28,
            boxstyle="round,pad=0.03",
            facecolor=bg, edgecolor=edge, linewidth=1.5, zorder=4
        )
        ax.add_patch(bubble)
        ax.text(x, y, label, ha="center", va="center", fontsize=8.5,
                color=txt, fontweight="bold", zorder=5)

    ax.set_title("統合失調症の発症に関わる主な要因", fontsize=10,
                 color="#1E3A5F", fontweight="bold", pad=10)
    plt.tight_layout(pad=0.5)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    return buf


# ---------------------------------------------------------------------------
# 各章の docx 構築
# ---------------------------------------------------------------------------

def chapter1(doc: Document, img_age: io.BytesIO) -> None:
    add_h1(doc, "① 統合失調症ってどんな病気？")

    add_para(doc, "統合失調症は、脳のはたらき方が一時的に変わることで、「ものの聞こえ方・見え方」「考え方」「感じ方」に普段とは違うことが起きる病気です。性格の問題でも本人の努力不足でもなく、脳の医学的な病気です。")

    # 統計ボックス3枚（3列テーブル）
    add_h2(doc, "知っておきたい数字・事実")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    stat_data = [
        ("約 1 %", "世界中で100人に1人\n—クラス40人なら身近にいる可能性", C_STAT1_BG, "3B82F6"),
        ("10代後半〜", "30代に多い\n—脳が大きく変化する時期", C_STAT2_BG, "0F766E"),
        ("2002年", "「精神分裂病」→「統合失調症」\n—偏見をなくすために改称", C_STAT3_BG, "6366F1"),
    ]
    cells = tbl.rows[0].cells
    for i, (num, desc, bg, col) in enumerate(stat_data):
        cell = cells[i]
        set_cell_shading(cell, bg)
        set_cell_padding(cell, 120)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(2)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p1.paragraph_format.line_spacing = 1.2
        add_run(p1, num, bold=True, size=Pt(15), color_hex=col)
        p2 = cell.add_paragraph()
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p2.paragraph_format.line_spacing = 1.25
        add_run(p2, desc, size=PT_SMALL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 年代別棒グラフ
    add_h2(doc, "発症しやすい年代（図1）")
    insert_image(doc, img_age, width_cm=13.5)
    add_para(doc, "※ 相対的なイメージ図であり、正確な統計値ではありません。",
             size=PT_CAPTION, color_hex="64748B")

    # infoボックス
    add_colored_box(doc,
        "ポイント",
        "適切な治療と支援があれば、多くの人が症状をコントロールしながら学校・仕事・日常生活を送れるようになります。「一生治らない」は誤解です。",
        C_BLUE_LIGHT, "3B82F6", "1D4ED8")


def chapter2(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "② どんなことが起きる？（症状）")

    add_para(doc, "症状は大きく3つのグループに分けると整理しやすいです。すべての人に同じ症状が出るわけではなく、組み合わせや強さは人それぞれです。")

    # 3列カラー表
    tbl = doc.add_table(rows=5, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = [
        ("＋ 陽性症状", C_RED_MID, "7F1D1D", '"普段はないものが加わる"'),
        ("－ 陰性症状", C_BLUE_MID, "1E3A8A", '"普段あるものが減る"'),
        ("🧠 認知症状", C_PURPLE_MID, "4C1D95", '"考える力に影響が出る"'),
    ]
    rows_data = [
        [("幻聴", "誰もいないのに声が聞こえる"),
         ("意欲の低下", "何もやる気が出ない"),
         ("集中力の低下", "授業・会話に集中できない")],
        [("妄想", "「監視されている」など証拠なく強く信じる"),
         ("感情が薄れる", "喜怒哀楽が感じにくい・表情が乏しい"),
         ("記憶の問題", "さっきのことを覚えにくい")],
        [("思考の混乱", "考えがバラバラになる"),
         ("社会的引きこもり", "人との関わりが億劫になる"),
         ("計画が立てにくい", "段取りよく行動しにくくなる")],
    ]
    bg_cols = [C_RED_LIGHT, C_BLUE_LIGHT, C_PURPLE_LIGHT]
    hdr_cols = [C_RED_MID, C_BLUE_MID, C_PURPLE_MID]

    # ヘッダー行
    for j, (title, bg, txt_col, subtitle) in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        set_cell_shading(cell, bg)
        set_cell_padding(cell, 110)
        cell.text = ""
        ph = cell.paragraphs[0]
        ph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_after = Pt(1)
        ph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        ph.paragraph_format.line_spacing = 1.2
        add_run(ph, title, bold=True, size=Pt(11), color_hex=txt_col)
        ps = cell.add_paragraph()
        ps.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ps.paragraph_format.space_after = Pt(0)
        ps.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        ps.paragraph_format.line_spacing = 1.2
        add_run(ps, subtitle, size=PT_SMALL, italic=True, color_hex=txt_col)

    # データ行
    for ri, row_items in enumerate(rows_data):
        for j, (term, desc) in enumerate(row_items):
            cell = tbl.rows[ri + 1].cells[j]
            set_cell_shading(cell, bg_cols[j])
            set_cell_padding(cell, 110)
            cell.text = ""
            pt = cell.paragraphs[0]
            pt.paragraph_format.space_after = Pt(1)
            pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pt.paragraph_format.line_spacing = 1.2
            add_run(pt, f"▶ {term}", bold=True, size=PT_BODY)
            pd = cell.add_paragraph()
            pd.paragraph_format.space_after = Pt(0)
            pd.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pd.paragraph_format.line_spacing = 1.25
            add_run(pd, desc, size=PT_SMALL, color_hex="475569")

    # 空白行 (4行目以降を使わないので詰める)
    for ri in [4]:
        for j in range(3):
            cell = tbl.rows[ri].cells[j]
            set_cell_shading(cell, "FFFFFF")
            cell.text = ""

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_colored_box(doc,
        "症状には「波」があります",
        "症状が強い急性期と、落ち着いている安定期を繰り返すことが多いです。調子が良い時期は、症状がほとんどわからないこともあります。",
        C_AMBER_LIGHT, "F59E0B", "B45309")


def chapter3(doc: Document, img_cause: io.BytesIO) -> None:
    doc.add_page_break()
    add_h1(doc, "③ なぜなるの？（原因）")

    add_para(doc, "「これが原因！」と断言できるものはまだ一つもありません。現在の研究では、いくつかの要因が複合的に重なって発症すると考えられています。")

    # 放射図
    add_h2(doc, "発症に関わる要因（図2）")
    insert_image(doc, img_cause, width_cm=13.0)

    # 要因テーブル
    add_h2(doc, "各要因のくわしい説明")
    factors = [
        ("脳の神経伝達物質", "脳の中で情報を伝える「ドーパミン」などのバランスが乱れることが関わっていると考えられている"),
        ("遺伝的要因", "家族に同じ病気の人がいると、やや発症しやすくなる。ただし、家族が発症しても必ずなるわけではない"),
        ("強いストレス・睡眠不足", "長期のストレスや睡眠不足、大麻などの薬物使用が発症の引き金になることがある"),
        ("環境・脳の発達", "胎児のころの脳の発達や、思春期の脳の変化が関係している可能性がある"),
    ]
    tbl = doc.add_table(rows=len(factors) + 1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell_para(tbl.rows[0].cells[0], "要因", bold=True, size=PT_H2, color_hex="FFFFFF")
    set_cell_shading(tbl.rows[0].cells[0], C_PRIMARY)
    set_cell_padding(tbl.rows[0].cells[0], 110)
    cell_para(tbl.rows[0].cells[1], "説明", bold=True, size=PT_H2, color_hex="FFFFFF")
    set_cell_shading(tbl.rows[0].cells[1], C_PRIMARY)
    set_cell_padding(tbl.rows[0].cells[1], 110)

    for i, (name, desc) in enumerate(factors):
        bg = C_ROW_ALT if i % 2 == 0 else "FFFFFF"
        cell_para(tbl.rows[i + 1].cells[0], name, bold=True, size=PT_BODY)
        set_cell_shading(tbl.rows[i + 1].cells[0], bg)
        set_cell_padding(tbl.rows[i + 1].cells[0], 110)
        cell_para(tbl.rows[i + 1].cells[1], desc, size=PT_BODY)
        set_cell_shading(tbl.rows[i + 1].cells[1], bg)
        set_cell_padding(tbl.rows[i + 1].cells[1], 110)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_colored_box(doc,
        "大切なこと",
        "「育て方が悪かった」「本人が弱い」は間違いです。生物学的・環境的な複数の要因が重なって起こる病気であり、本人や家族のせいではありません。",
        C_AMBER_LIGHT, "F59E0B", "B45309")


def chapter4(doc: Document, img_recovery: io.BytesIO) -> None:
    doc.add_page_break()
    add_h1(doc, "④ 治療と回復について")

    add_para(doc, "統合失調症は、適切な治療と支援を続けることで、多くの人が症状をコントロールしながら生活できるようになります。")

    # 治療3ステップ表
    add_h2(doc, "主な治療の3本柱")
    steps = [
        ("Step 1", "薬物療法", "脳の神経伝達物質のバランスを整える「抗精神病薬」が中心。飲み続けることで再発を防ぎやすくなる。★ 勝手にやめないことが大切。", C_STAT2_BG, "0F766E"),
        ("Step 2", "心理教育・カウンセリング", "病気・薬のことを正しく学ぶ。つらい体験（幻聴など）への対処法を一緒に考える「認知行動療法」も有効。", C_BLUE_LIGHT, "3B82F6"),
        ("Step 3", "リハビリ・生活支援", "デイケアで生活リズムを整えたり、仲間と活動したり、就労の準備をする。回復の目標は本人が決める。", C_PURPLE_LIGHT, "7C3AED"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = tbl.rows[0].cells
    for i, (step, title, desc, bg, col) in enumerate(steps):
        cell = cells[i]
        set_cell_shading(cell, bg)
        set_cell_padding(cell, 110)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cell.text = ""
        ps = cell.paragraphs[0]
        ps.paragraph_format.space_after = Pt(1)
        ps.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        ps.paragraph_format.line_spacing = 1.2
        add_run(ps, step, bold=True, size=PT_SMALL, color_hex=col)
        pt = cell.add_paragraph()
        pt.paragraph_format.space_after = Pt(3)
        pt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pt.paragraph_format.line_spacing = 1.2
        add_run(pt, title, bold=True, size=PT_H2, color_hex=col)
        pd = cell.add_paragraph()
        pd.paragraph_format.space_after = Pt(0)
        pd.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pd.paragraph_format.line_spacing = LINE_SPACING
        add_run(pd, desc, size=PT_SMALL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 回復曲線
    add_h2(doc, "回復のイメージ（図3）")
    insert_image(doc, img_recovery, width_cm=13.5)
    add_para(doc, "回復は一直線ではなく、波を繰り返しながら全体的に改善していきます。",
             size=PT_CAPTION, color_hex="64748B")

    add_colored_box(doc,
        "早期発見・早期治療が大切",
        "早く気づいて治療を始めるほど回復が早い傾向があります。「症状がゼロになること」だけが回復ではなく、学校・仕事・趣味・人間関係など、自分が大切にしたいことができるようになることが目標です。",
        C_GREEN_LIGHT, "22C55E", "15803D")


def chapter5(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "⑤ 日常生活・学校・仕事への影響")

    add_para(doc, "支援をうまく使うことで、多くの人が学校や仕事を続けたり再開したりできます。")

    # 場面別3列表
    add_h2(doc, "よくある困りごとと対策")
    scenes = [
        ("学校", "集中が続かない・休みがちになる",
         "・スクールカウンセラーに相談\n・休学・復学の制度を利用\n・通信制・定時制も選択肢",
         C_STAT1_BG, "3B82F6"),
        ("仕事", "疲れやすい・ミスが増える",
         "・障害者雇用で配慮を受ける\n・就労移行支援を利用\n・短時間勤務から始める",
         C_STAT2_BG, "0F766E"),
        ("生活費・お金", "働けない時期の収入が心配",
         "・障害年金を申請\n・自立支援医療で通院費を削減\n・精神障害者手帳でサービス利用",
         C_STAT3_BG, "6366F1"),
    ]
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    row_labels = ["場面", "よくある困りごと", "できる対策・工夫"]
    for ri, label in enumerate(row_labels):
        for ci, (title, issue, actions, bg, col) in enumerate(scenes):
            cell = tbl.rows[ri].cells[ci]
            set_cell_padding(cell, 110)
            if ri == 0:
                set_cell_shading(cell, bg)
                cell_para(cell, title, bold=True, size=PT_H2, color_hex=col,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            elif ri == 1:
                set_cell_shading(cell, "FFFBEB")
                cell_para(cell, issue, size=PT_BODY)
            else:
                set_cell_shading(cell, C_GREEN_LIGHT)
                cell_para(cell, actions, size=PT_SMALL)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 社会制度一覧
    add_h2(doc, "使える主な社会的支援制度")
    institutions = [
        ("自立支援医療（精神通院）", "通院の医療費が原則1割に。申請は市区町村の窓口へ。"),
        ("精神障害者保健福祉手帳", "交通費割引・税制優遇・就職支援などが受けられる。"),
        ("障害年金", "病気で働きにくい場合に受け取れる年金（1〜2級）。"),
        ("就労移行支援", "働く準備を最大2年間サポート。原則無料で利用可。"),
        ("精神保健福祉センター", "各都道府県に設置。本人・家族・周囲の人も無料で相談可。"),
    ]
    tbl2 = doc.add_table(rows=len(institutions) + 1, cols=2)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell_para(tbl2.rows[0].cells[0], "制度名", bold=True, size=PT_H2, color_hex="FFFFFF")
    set_cell_shading(tbl2.rows[0].cells[0], C_TEAL)
    set_cell_padding(tbl2.rows[0].cells[0], 110)
    cell_para(tbl2.rows[0].cells[1], "内容", bold=True, size=PT_H2, color_hex="FFFFFF")
    set_cell_shading(tbl2.rows[0].cells[1], C_TEAL)
    set_cell_padding(tbl2.rows[0].cells[1], 110)

    for i, (name, desc) in enumerate(institutions):
        bg = C_ROW_ALT if i % 2 == 0 else "FFFFFF"
        cell_para(tbl2.rows[i + 1].cells[0], name, bold=True, size=PT_BODY)
        set_cell_shading(tbl2.rows[i + 1].cells[0], bg)
        set_cell_padding(tbl2.rows[i + 1].cells[0], 110)
        cell_para(tbl2.rows[i + 1].cells[1], desc, size=PT_BODY)
        set_cell_shading(tbl2.rows[i + 1].cells[1], bg)
        set_cell_padding(tbl2.rows[i + 1].cells[1], 110)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def chapter6(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "⑥ 偏見をなくすために／サポートの仕方")

    # 誤解 vs 正解 2列テーブル
    add_h2(doc, "よくある誤解と、正しい理解")
    myths = [
        ("「危険な人だ」", "他者を傷つけるリスクは一般の人と大きく変わらない。孤立や自傷への配慮が必要なことが多い。"),
        ("「一生治らない」", "適切な治療で症状をコントロールできる。多くの人が仕事・学業・生活を取り戻す。"),
        ("「話が通じない」", "急性期（症状が強い時期）以外は、普通に会話できることがほとんど。"),
        ("「自業自得・意志が弱い」", "脳の医学的な病気。本人の責任ではない。"),
        ("「精神病=怖い」", "正しい知識が偏見をなくす。誤解や無関心が本人を苦しめる。"),
    ]
    tbl = doc.add_table(rows=len(myths) + 1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell_para(tbl.rows[0].cells[0], "× よくある誤解", bold=True, size=PT_H2, color_hex="FFFFFF",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(tbl.rows[0].cells[0], "DC2626")
    set_cell_padding(tbl.rows[0].cells[0], 110)
    cell_para(tbl.rows[0].cells[1], "✓ 正しい理解", bold=True, size=PT_H2, color_hex="FFFFFF",
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(tbl.rows[0].cells[1], "15803D")
    set_cell_padding(tbl.rows[0].cells[1], 110)

    for i, (myth, truth) in enumerate(myths):
        bg_l = "FFF1F2" if i % 2 == 0 else "FEE2E2"
        bg_r = "F0FDF4" if i % 2 == 0 else "DCFCE7"
        cell_para(tbl.rows[i + 1].cells[0], myth, size=PT_BODY, color_hex="991B1B")
        set_cell_shading(tbl.rows[i + 1].cells[0], bg_l)
        set_cell_padding(tbl.rows[i + 1].cells[0], 110)
        cell_para(tbl.rows[i + 1].cells[1], truth, size=PT_BODY, color_hex="166534")
        set_cell_shading(tbl.rows[i + 1].cells[1], bg_r)
        set_cell_padding(tbl.rows[i + 1].cells[1], 110)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # サポート5か条
    add_h2(doc, "困っている人へのサポート 5か条")
    tips = [
        ("1", "まず「話を聞く」", "解決しようとせず「大変だったね」と気持ちを受け止めるだけでいい", "0F766E", C_STAT2_BG),
        ("2", "否定・説得しない", "「気のせいだよ」「頑張れば大丈夫」は逆効果になることがある", "3B82F6", C_STAT1_BG),
        ("3", "専門家につなぐ", "「一緒に相談窓口に行こう」と誘う。一人で抱え込まなくていい", "6366F1", C_STAT3_BG),
        ("4", "秘密を守る", "聞いたことを本人の許可なく周囲に広めない", "7C3AED", C_PURPLE_LIGHT),
        ("5", "自分も無理しない", "サポートする側も疲弊しないよう、自分のペースで関わる", "E11D48", C_RED_LIGHT),
    ]
    tbl2 = doc.add_table(rows=len(tips), cols=3)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (num, title, desc, col, bg) in enumerate(tips):
        nc = tbl2.rows[i].cells[0]
        set_cell_shading(nc, col)
        set_cell_padding(nc, 80)
        nc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_para(nc, num, bold=True, size=Pt(14), color_hex="FFFFFF",
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        tc = tbl2.rows[i].cells[1]
        set_cell_shading(tc, bg)
        set_cell_padding(tc, 110)
        tc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_para(tc, title, bold=True, size=PT_H2, color_hex=col)

        dc = tbl2.rows[i].cells[2]
        set_cell_shading(dc, "FFFFFF")
        set_cell_padding(dc, 110)
        dc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_para(dc, desc, size=PT_BODY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 相談窓口
    add_h2(doc, "相談できる主な窓口")
    contacts = [
        ("精神保健福祉センター", "各都道府県に設置。本人・家族・周囲の人も無料で相談可。", "DBEAFE", "1D4ED8"),
        ("こころの健康相談統一ダイヤル", "0570-064-556（全国共通・平日）", "DCFCE7", "15803D"),
        ("よりそいホットライン", "0120-279-338（24時間・無料）", "FEE2E2", "DC2626"),
        ("スクールカウンセラー", "在学中は学校内で相談可能", "F3E8FF", "7C3AED"),
    ]
    tbl3 = doc.add_table(rows=len(contacts), cols=2)
    tbl3.style = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (name, detail, bg, col) in enumerate(contacts):
        nc = tbl3.rows[i].cells[0]
        set_cell_shading(nc, bg)
        set_cell_padding(nc, 110)
        nc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_para(nc, name, bold=True, size=PT_BODY, color_hex=col)
        dc = tbl3.rows[i].cells[1]
        set_cell_shading(dc, "FFFFFF")
        set_cell_padding(dc, 110)
        dc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_para(dc, detail, size=PT_BODY)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_colored_box(doc,
        "この資料を読んで",
        "統合失調症は、正しく理解すれば「怖い病気」ではなく、「適切な支援があれば回復できる病気」です。偏見や無関心ではなく、正しい知識と温かい関わりが、本人を助け社会全体をより生きやすくします。",
        C_BLUE_LIGHT, "3B82F6", "1D4ED8")


def footer_page(doc: Document) -> None:
    doc.add_page_break()
    add_h1(doc, "（付記）")
    notes = [
        "この資料は、統合失調症について正しく知ってもらうことを目的とした一般教育用の資料です。",
        "個別の診断・治療については、必ず医療機関にご相談ください。",
        "参考: 日本精神神経学会、厚生労働省「みんなのメンタルヘルス」、世界保健機関（WHO）（2026年5月）",
    ]
    for note in notes:
        add_para(doc, note, size=PT_SMALL, color_hex="475569")


# ---------------------------------------------------------------------------
# メインドキュメント構築
# ---------------------------------------------------------------------------

def build_document(out_path: Path) -> None:
    print("図を生成中...")
    img_age = fig_age_bar()
    img_cause = fig_cause_radial()
    img_recovery = fig_recovery_curve()

    print("Word文書を構築中...")
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)

    # タイトル
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "統合失調症って何？", bold=True, size=PT_TITLE, color_hex=C_PRIMARY)

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(2)
    add_run(p2, "高校生・一般向け 図解ガイド", size=Pt(11.5), color_hex=C_TEAL)

    p3 = doc.add_paragraph()
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_after = Pt(10)
    add_run(p3, "作成日: 2026年5月", size=PT_SMALL, color_hex="94A3B8")

    chapter1(doc, img_age)
    chapter2(doc)
    chapter3(doc, img_cause)
    chapter4(doc, img_recovery)
    chapter5(doc)
    chapter6(doc)
    footer_page(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="統合失調症 高校生向け図解版 .docx 生成"
    )
    parser.add_argument(
        "--output", "-o",
        default="samples/output/統合失調症_高校生向け解説_図解版.docx",
        help="出力ファイルパス"
    )
    ns = parser.parse_args(argv)
    build_document(Path(ns.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
