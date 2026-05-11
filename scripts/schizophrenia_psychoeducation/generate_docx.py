#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
統合失調症の心理教育資料（.docx）をオフラインで生成する。

使用例:
  python generate_docx.py --input example_case.json --output 心理教育_本人.docx
  python generate_docx.py --input example_case.json --audience family --output 心理教育_家族.docx

content.json は章ごとに blocks 配列（type: p | h2 | bullets | table）を記述できる。
レガシー paragraphs（文字列配列のみ）にも後方互換。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT_DIR = Path(__file__).resolve().parent
CONTENT_PATH = ROOT_DIR / "content.json"

# 院内・印刷見やすさ優先。色使いはごく薄いグレーのみ（モノクロ印刷でも判別可）。
BODY_FONT = "MS ゴシック"
HEADING_FONT = "MS ゴシック"
TITLE_FONT = "MS ゴシック"
TABLE_STYLE = "Table Grid"

# 視認性: 本文は11pt、行間は約1.35倍。見出しは段階をはっきり。
PT_BODY = Pt(11)
PT_TITLE = Pt(15)
PT_H1 = Pt(13)
PT_H2 = Pt(12)
PT_CAPTION = Pt(10.5)
PT_SUBTITLE = Pt(11.5)

LINE_SPACING_MULTIPLE = 1.35

# 表: ヘッダー淡色、データ行は交互の白／極薄グレー
COLOR_TABLE_HEADER = "E4E6EB"
COLOR_TABLE_ROW_ALT = "F3F4F7"
COLOR_TABLE_ROW_BASE = "FFFFFF"
COLOR_RULE_H1 = "B0B0B0"

CELL_PADDING_DXA = 115  # セル内余白（約2mm）

PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

RESERVED_CASE_KEYS = frozenset({"audience", "meta", "content_version"})

# 新版プレースホルダ（章立て案A・6章）。旧キーは CASE_KEY_ALIASES で補完。
DEFAULT_FIELDS: dict[str, str] = {
    "facility_name": "当院",
    "ch1_case_note": "（この症例の症状の要点を、できるだけ本人／家族にも分かる言葉で記入してください。）",
    "ch2_known": "（説明できる医学的事実や、この症例に当てはまりうるしくみを記入してください。）",
    "ch2_unknown": "（断言が難しい点／検証中／個人差が大きい点を記入してください。）",
    "ch3_short": "（数週〜数か月の方向性を記入してください。）",
    "ch3_long": "（数か月〜年単位の方向性を記入してください。）",
    "ch4_medication": "（一般名・用法・頓服・モニタリング・主治医への相談目安などを記入してください。）",
    "ch5_nonpharma": "（心理社会的支援・生活調整・家族の関わり方などを記入してください。）",
    "ch6_early_warning": "（早期警告サインを、本人語／家族視点など分けてでもよいので記入してください。）",
    "ch6_trouble_contact": "（困ったときの連絡先・受診基準・夜間・祝日などを記入してください。）",
    "ch6_support": "（利用検討中の制度・相談窓口・手続き状況を記入してください。）",
}

# 症例 JSON の旧キー → 新版キー（新版が空／未指定のときだけ転写）
CASE_KEY_ALIASES: tuple[tuple[str, str], ...] = (
    ("ch2_known", "ch1_known"),
    ("ch2_unknown", "ch1_unknown"),
    ("ch3_short", "ch2_short"),
    ("ch3_long", "ch2_long"),
    ("ch4_medication", "ch3_medication"),
    ("ch5_nonpharma", "ch4_nonpharma"),
    ("ch6_support", "ch5_support"),
)


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def set_paragraph_body_rhythm(paragraph, *, tight: bool = False) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING_MULTIPLE
    pf.widow_control = True
    pf.space_after = Pt(4 if tight else 6)


def set_paragraph_h1_border(paragraph) -> None:
    """章見出し下に細いラインを入れ、横断スキャンを助ける（色は単色印刷でも破綻しにくい）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), COLOR_RULE_H1)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_padding(cell, dxa: int = CELL_PADDING_DXA) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    old_mar = tc_pr.find(qn("w:tcMar"))
    if old_mar is not None:
        tc_pr.remove(old_mar)
    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(dxa))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def format_table_cell(cell, *, valign: WD_CELL_VERTICAL_ALIGNMENT | None = None) -> None:
    set_cell_padding(cell)
    if valign is not None:
        cell.vertical_alignment = valign


def add_title_line(doc: Document, text: str, *, size_pt: Pt) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, TITLE_FONT)
    r.font.size = size_pt
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.widow_control = True


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, HEADING_FONT)
    r.font.bold = True
    r.font.size = PT_H1
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True
    set_paragraph_h1_border(p)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_east_asia_font(r, HEADING_FONT)
    r.font.bold = True
    r.font.size = PT_H2
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True


def add_para_line(
    doc: Document,
    line: str,
    *,
    size_pt: Pt = PT_BODY,
    bold: bool = False,
    style: str | None = None,
    tight: bool = False,
) -> None:
    kw: dict[str, Any] = {}
    if style:
        kw["style"] = style
    p = doc.add_paragraph(**kw)
    r = p.add_run(line)
    set_east_asia_font(r, BODY_FONT)
    r.font.size = size_pt
    r.font.bold = bold
    set_paragraph_body_rhythm(p, tight=tight)


def add_paragraph_block(doc: Document, text: str, *, size_pt: Pt = PT_BODY) -> None:
    for line in text.split("\n"):
        s = line.strip()
        if s:
            add_para_line(doc, s, size_pt=size_pt)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        s = item.strip()
        if not s:
            continue
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(s)
        set_east_asia_font(r, BODY_FONT)
        r.font.size = PT_BODY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = LINE_SPACING_MULTIPLE
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-6)
        p.paragraph_format.widow_control = True


def add_table_block(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    caption: str | None = None,
    mapping: Mapping[str, str],
    missing: list[str],
) -> None:
    if caption:
        add_para_line(
            doc,
            substitute(caption, mapping, missing),
            size_pt=PT_CAPTION,
            bold=True,
            tight=True,
        )

    ncols = len(headers)
    nrows = 1 + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        cell = hdr[j]
        set_cell_shading(cell, COLOR_TABLE_HEADER)
        format_table_cell(cell, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(substitute(str(h), mapping, missing))
        set_east_asia_font(r, BODY_FONT)
        r.font.bold = True
        r.font.size = PT_BODY
        set_paragraph_table_cell(p)

    for i, row in enumerate(rows):
        stripe = COLOR_TABLE_ROW_ALT if i % 2 == 0 else COLOR_TABLE_ROW_BASE
        for j in range(ncols):
            cell = table.rows[i + 1].cells[j]
            val = row[j] if j < len(row) else ""
            set_cell_shading(cell, stripe)
            format_table_cell(cell, valign=WD_CELL_VERTICAL_ALIGNMENT.TOP)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(substitute(str(val), mapping, missing))
            set_east_asia_font(r, BODY_FONT)
            r.font.size = PT_BODY
            set_paragraph_table_cell(p)

    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(10)


def set_paragraph_table_cell(paragraph) -> None:
    """表セル内: 読みやすい行間、余計な段落間隔を詰める。"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING_MULTIPLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.widow_control = True


def substitute(text: str, mapping: Mapping[str, str], missing: list[str]) -> str:
    def repl(m: re.Match[str]) -> str:
        key = m.group(1)
        if key not in mapping:
            missing.append(key)
            return ""
        return str(mapping[key])

    return PLACEHOLDER_RE.sub(repl, text)


def load_json(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as f:
        return json.load(f)


def build_placeholder_map(case: Mapping[str, Any]) -> dict[str, str]:
    m: dict[str, str] = {k: v for k, v in DEFAULT_FIELDS.items()}

    for k, v in case.items():
        if k in RESERVED_CASE_KEYS:
            continue
        if k == "placeholders" and isinstance(v, dict):
            for pk, pv in v.items():
                m[str(pk)] = "" if pv is None else str(pv)
        else:
            m[str(k)] = "" if v is None else str(v)

    def nonempty(v: Any) -> bool:
        return v is not None and str(v).strip() != ""

    for new_k, old_k in CASE_KEY_ALIASES:
        if nonempty(case.get(old_k)) and not nonempty(case.get(new_k)):
            m[new_k] = str(case[old_k])

    return m


def blocks_for_chapter(chapter: Mapping[str, Any], audience: str) -> list[Any]:
    if "blocks" in chapter:
        return list(chapter["blocks"].get(audience, []))
    raw = chapter.get("paragraphs", {}).get(audience, [])
    return [{"type": "p", "text": str(x)} for x in raw]


def render_block(
    doc: Document,
    block: Mapping[str, Any] | str,
    *,
    audience: str,
    mapping: Mapping[str, str],
    missing: list[str],
) -> None:
    if isinstance(block, str):
        add_paragraph_block(doc, substitute(block, mapping, missing))
        return

    btype = block.get("type", "p")

    if btype in ("p", "paragraph"):
        text = substitute(str(block.get("text", "")), mapping, missing)
        add_paragraph_block(doc, text)
        return

    if btype == "h2":
        add_h2(doc, substitute(str(block.get("text", "")), mapping, missing))
        return

    if btype == "bullets":
        items = block.get("items", [])
        filled = [substitute(str(x), mapping, missing) for x in items]
        add_bullets(doc, filled)
        return

    if btype == "table":
        headers = [str(x) for x in block.get("headers", [])]
        rows_raw = block.get("rows", [])
        rows: list[list[str]] = []
        for row in rows_raw:
            rows.append([str(c) for c in row])
        caption = block.get("caption")
        cap = str(caption) if caption else None
        add_table_block(doc, headers, rows, caption=cap, mapping=mapping, missing=missing)
        return

    if btype == "both":
        inner = block.get(audience) or block.get("patient")
        if isinstance(inner, list):
            for sub in inner:
                render_block(doc, sub, audience=audience, mapping=mapping, missing=missing)
        elif isinstance(inner, dict):
            render_block(doc, inner, audience=audience, mapping=mapping, missing=missing)
        elif isinstance(inner, str):
            add_paragraph_block(doc, substitute(inner, mapping, missing))
        return

    # 未知タイプはテキストとして
    if "text" in block:
        add_paragraph_block(doc, substitute(str(block["text"]), mapping, missing))


def validate_audience(a: str) -> str:
    a = a.strip().lower()
    if a in {"patient", "本人", "honnin"}:
        return "patient"
    if a in {"family", "家族", "kazoku"}:
        return "family"
    raise ValueError(f"audience は patient か family としてください: {a!r}")


def generate_docx(
    *,
    content: Mapping[str, Any],
    case: Mapping[str, Any],
    audience: str,
) -> Document:
    audience = validate_audience(audience)
    mapping = build_placeholder_map(case)
    missing: list[str] = []

    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Cm(2.15)
    sec.bottom_margin = Cm(2.15)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.15)

    title = content["titles"][audience]
    add_title_line(doc, title, size_pt=PT_TITLE)

    if case.get("subject_line"):
        add_title_line(doc, str(case["subject_line"]), size_pt=PT_SUBTITLE)
    if case.get("date_line"):
        add_title_line(doc, str(case["date_line"]), size_pt=PT_CAPTION)

    doc.add_paragraph()

    for idx, chapter in enumerate(content["chapters"]):
        if idx > 0:
            doc.add_page_break()

        heading = chapter["heading"]
        h1_text = substitute(str(heading), mapping, missing)
        add_h1(doc, h1_text)

        for block in blocks_for_chapter(chapter, audience):
            if isinstance(block, (dict, str)):
                render_block(doc, block, audience=audience, mapping=mapping, missing=missing)
            else:
                render_block(
                    doc,
                    {"type": "p", "text": str(block)},
                    audience=audience,
                    mapping=mapping,
                    missing=missing,
                )

    doc.add_page_break()
    add_h1(doc, "（付記）")
    for line in content.get("footer", []):
        add_paragraph_block(doc, substitute(str(line), mapping, missing))

    leaked = [p.text[:80] for p in doc.paragraphs if "{{" in p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "{{" in cell.text:
                    leaked.append(cell.text[:80])

    if missing:
        print(
            "警告: 症例データに存在しないプレースホルダキーがあり、空に置き換えました: "
            + ", ".join(sorted(set(missing))),
            file=sys.stderr,
        )
    if leaked:
        print(
            "警告: 文面に {{...}} が残っています: " + str(leaked[:5]),
            file=sys.stderr,
        )

    return doc


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="統合失調症 心理教育資料 .docx 生成（オフライン）",
    )
    parser.add_argument("--input", "-i", required=True, help="症例 JSON")
    parser.add_argument("--output", "-o", required=True, help="出力 .docx")
    parser.add_argument(
        "--audience",
        "-a",
        default=None,
        help="patient または family。省略時は JSON の audience",
    )
    parser.add_argument(
        "--content",
        "-c",
        default=str(CONTENT_PATH),
        help="定型文 JSON",
    )
    ns = parser.parse_args(argv)

    content_path = Path(ns.content)
    case_path = Path(ns.input)
    out_path = Path(ns.output)

    content = load_json(content_path)
    case = load_json(case_path)

    aud = ns.audience if ns.audience else case.get("audience")
    if not aud:
        print("症例 JSON に audience または --audience を指定してください。", file=sys.stderr)
        return 2

    try:
        doc = generate_docx(content=content, case=case, audience=str(aud))
    except ValueError as e:
        print(str(e), file=sys.stderr)
        return 2

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
