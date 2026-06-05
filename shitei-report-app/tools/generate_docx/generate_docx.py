#!/usr/bin/env python3
"""
精神保健指定医 ケースレポート .docx 生成スクリプト
様式3-1（令和7年1月版）準拠

使用方法:
    python3 generate_docx.py --input sample_case.json
    python3 generate_docx.py --input_dir ./cases/
    python3 generate_docx.py --all_samples  # サンプルデータで全5症例を生成

依存パッケージ:
    pip3 install python-docx
"""

import json
import os
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx がインストールされていません。")
    print("  pip3 install python-docx")
    sys.exit(1)


# ──────────────────────────────────────────────
# 定数
# ──────────────────────────────────────────────
CASE_CATEGORY_LABELS = {
    1: "第1症例：症状性を含む器質性精神障害（F0）",
    2: "第2症例：精神作用物質依存症（F1 依存症）",
    3: "第3症例：統合失調症、統合失調症型障害及び妄想性障害（F2）",
    4: "第4症例：気分（感情）障害（F3）",
    5: "第5症例：F4〜F9のいずれか（神経症性障害・パーソナリティ障害・発達障害等）",
}

ADMISSION_TYPES = {
    "措置入院": "措置入院（精神保健福祉法 第29条）",
    "医療保護入院": "医療保護入院（精神保健福祉法 第33条）",
}


# ──────────────────────────────────────────────
# スタイル設定ヘルパー
# ──────────────────────────────────────────────
def set_font(run, size_pt=10.5, bold=False, color_rgb=None, font_name="ＭＳ 明朝"):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    # East Asian font
    r = run._r
    rpr = r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), font_name)
    rpr.insert(0, rfonts)


def add_heading(doc, text, level=1):
    """節・項のヘッダーを追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size_pt=12 if level == 1 else 10.5, bold=True,
             color_rgb=(30, 58, 95))
    return p


def add_field(doc, label, value, indent=False):
    """ラベル: 値 の行を追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run_label = p.add_run(f"【{label}】")
    set_font(run_label, size_pt=9.5, bold=True, color_rgb=(15, 118, 110))
    run_val = p.add_run("　" + (value or "（記入なし）"))
    set_font(run_val, size_pt=10)
    return p


def add_body_text(doc, text):
    """本文テキストを追加"""
    p = doc.add_paragraph(text or "（記入なし）")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_font(run, size_pt=10)
    return p


def add_separator(doc):
    """区切り線"""
    p = doc.add_paragraph("─" * 35)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        set_font(run, size_pt=8, color_rgb=(156, 163, 175))


def add_note_box(doc, text):
    """注意ボックス風テキスト"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"※ {text}")
    set_font(run, size_pt=9, color_rgb=(120, 53, 15))


# ──────────────────────────────────────────────
# メイン生成関数
# ──────────────────────────────────────────────
def generate_docx(case_data: dict, output_path: str):
    """
    case_data の構造（すべてのキーはオプション）:
      case_number          int   1〜5
      application_date     str   "2025-04-01"
      admission_type       str   "措置入院" | "医療保護入院"
      final_diagnosis      str   "統合失調症（妄想型）"
      icd_code             str   "F20.0"
      hospital             str   "○○病院 精神科"
      patient_initial      str   "T.H"
      patient_dob          str   "1978-06-15"
      admit_date           str   "2024-11-01"
      discharge_date       str   "2025-02-28"
      attend_start         str   "2024-11-01"
      supervisor           str   "田中一郎 / 指定医番号：12345"
      restriction          str   "有（隔離）" | "無"
      age                  int   45
      sex                  str   "男性"
      chief_complaint      str   "…"
      family_history       str   "…"
      life_history         str   "…"
      past_history         str   "…"
      pre_admission_course str   "…"
      admission_situation  str   "…"
      admission_course     str   "…"
      mental_status        dict  { behavior, consciousness, cognition, mood, thought,
                                   perception, risk, insight, other }
      diagnosis_basis      str   "…"
      diff_diagnosis       str   "…"
      admission_basis      str   "…"
      restriction_detail   str   "…"
      seisanin             str   "…"
      pharmacotherapy      str   "…"
      nonpharmacotherapy   str   "…"
      aftercare            str   "…"
      consideration        str   "…"
    """
    doc = Document()

    # ページ設定
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.2)

    case_num = case_data.get("case_number", 1)
    cat_label = CASE_CATEGORY_LABELS.get(case_num, f"第{case_num}症例")
    adm_type  = case_data.get("admission_type", "（入院形態未記入）")
    adm_full  = ADMISSION_TYPES.get(adm_type, adm_type)

    # ── タイトル ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ケースレポート")
    set_font(title_run, size_pt=16, bold=True, color_rgb=(30, 58, 95))

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(cat_label)
    set_font(sub_run, size_pt=12, bold=True, color_rgb=(15, 118, 110))

    note_p = doc.add_paragraph()
    note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_p.add_run("様式3-1（令和7年1月版）準拠")
    set_font(note_run, size_pt=9, color_rgb=(107, 114, 128))

    doc.add_paragraph()

    # ── 表紙情報 ──
    add_heading(doc, "【表紙情報】様式3-1 ①〜⑩", level=1)

    add_field(doc, "申請日（①記入日）",  case_data.get("application_date", ""))
    add_field(doc, "②最終診断名",        case_data.get("final_diagnosis", ""))
    add_field(doc, "  ICD-10コード",      case_data.get("icd_code", "")
              + "　（F以下2桁必須）")
    add_field(doc, "③主な評価対象の入院形態", adm_full
              + "\n　※ 任意入院は主評価対象に選択不可（様式3-1 ③）")
    add_field(doc, "④医療機関名",         case_data.get("hospital", ""))
    add_field(doc, "⑤患者イニシャル",     case_data.get("patient_initial", "")
              + "　生年月日：" + case_data.get("patient_dob", ""))
    add_field(doc, "⑥入院日",            case_data.get("admit_date", "")
              + "　退院日：" + case_data.get("discharge_date", ""))
    add_field(doc, "⑦-1 担当開始",        case_data.get("attend_start", ""))
    add_field(doc, "⑦-2 指導医",          case_data.get("supervisor", ""))
    add_note_box(doc, "⑦-3 指導医の自筆署名が様式に必須です（ソフト入力不可）。")
    add_field(doc, "⑩行動制限の有無",     case_data.get("restriction", "無"))

    add_separator(doc)

    # ── 本文 ──
    add_heading(doc, "【本 文】", level=1)

    ms = case_data.get("mental_status", {})
    adm_text = case_data.get("admission_situation", "")
    course_text = case_data.get("admission_course", "")
    consid_text = case_data.get("consideration", "")
    char_count = len(adm_text) + len(course_text) + len(consid_text)

    add_field(doc, "入院時診断名",  case_data.get("final_diagnosis", ""))
    add_field(doc, "最終診断名",    case_data.get("final_diagnosis", ""))
    add_field(doc, "患者の性別",    case_data.get("sex", ""))
    add_field(doc, "担当医になった時の患者の年齢", str(case_data.get("age", "")) + "歳")
    add_field(doc, "本文字数（入院時状況+入院後経過+考察）",
              f"{char_count}字　（規定：1200〜2500字）")

    doc.add_paragraph()

    add_heading(doc, "【初診時主訴】", level=2)
    add_body_text(doc, case_data.get("chief_complaint", ""))

    add_heading(doc, "【家族歴】", level=2)
    add_body_text(doc, case_data.get("family_history", ""))

    add_heading(doc, "【生育・生活歴】", level=2)
    add_body_text(doc, case_data.get("life_history", ""))

    add_heading(doc, "【既往歴】", level=2)
    add_body_text(doc, case_data.get("past_history", ""))

    add_heading(doc, "【現病歴】", level=2)

    add_heading(doc, "＜入院前経過＞", level=2)
    add_body_text(doc, case_data.get("pre_admission_course", ""))

    add_heading(doc, "＜入院時の状況＞　※字数カウント対象", level=2)
    p_note = doc.add_paragraph()
    run_note = p_note.add_run(
        f"[注：入院形態「{adm_type}」の入院必要性を法的要件に沿って記載。"
        "精神障害者であること・入院の必要性・"
        + ("自傷他害のおそれの根拠を具体的に記載。]"
           if adm_type == "措置入院"
           else "任意入院が行われる状態にない根拠（同意能力・病識欠如）を具体的に記載。]")
    )
    set_font(run_note, size_pt=8.5, color_rgb=(120, 53, 15))
    add_body_text(doc, adm_text)

    add_heading(doc, "＜入院後経過＞　※字数カウント対象", level=2)
    add_body_text(doc, course_text)

    # 精神医学的所見（参考）
    if any(ms.get(k) for k in ms):
        add_heading(doc, "【精神医学的所見（参考・本文に統合して記載）】", level=2)
        for key, label in [
            ("behavior", "外観・行動"),
            ("consciousness", "意識・見当識"),
            ("cognition", "知的機能・記憶"),
            ("mood", "気分・感情"),
            ("thought", "思考（形式・内容）"),
            ("perception", "知覚（幻覚等）"),
            ("risk", "自傷他害リスク"),
            ("insight", "病識・判断力"),
            ("other", "その他"),
        ]:
            val = ms.get(key, "")
            if val:
                add_field(doc, label, val, indent=True)

    add_separator(doc)

    # ── 診断・根拠 ──
    add_heading(doc, "【診断の根拠・入院の根拠】", level=1)

    add_heading(doc, "ICD-10診断基準との対応", level=2)
    add_body_text(doc, case_data.get("diagnosis_basis", ""))

    add_heading(doc, "鑑別診断と除外根拠", level=2)
    add_body_text(doc, case_data.get("diff_diagnosis", ""))

    add_heading(doc, f"{adm_type}の法的根拠・入院の必要性", level=2)
    add_body_text(doc, case_data.get("admission_basis", ""))

    add_separator(doc)

    # ── 法的手続 ──
    add_heading(doc, "【関係法規に定める手続への対応】", level=1)
    add_note_box(doc,
        "様式3-1の「関係法規に定める手続への対応」欄（チェックリスト）を"
        "印刷した様式に✓で記入してください。以下は補足記録です。")

    if case_data.get("restriction_detail"):
        add_heading(doc, "行動制限の詳細", level=2)
        add_body_text(doc, case_data.get("restriction_detail", ""))

    if case_data.get("seisanin"):
        add_heading(doc, "退院後生活環境相談員の選任（令和6年4月1日以降入院）", level=2)
        add_body_text(doc, case_data.get("seisanin", ""))

    add_separator(doc)

    # ── 治療・考察 ──
    add_heading(doc, "【治療計画・経過】", level=1)

    add_heading(doc, "薬物療法", level=2)
    add_body_text(doc, case_data.get("pharmacotherapy", ""))

    add_heading(doc, "非薬物療法・精神療法・リハビリテーション", level=2)
    add_body_text(doc, case_data.get("nonpharmacotherapy", ""))

    add_heading(doc, "退院後支援・通院治療", level=2)
    add_body_text(doc, case_data.get("aftercare", ""))

    if consid_text:
        add_heading(doc, "【考察】（任意・字数カウント対象）", level=2)
        add_body_text(doc, consid_text)

    add_separator(doc)

    # ── フッタ注意事項 ──
    footer_p = doc.add_paragraph()
    footer_run = footer_p.add_run(
        "【提出時の注意事項】\n"
        "・様式3-1（令和7年1月版）を使用すること\n"
        "・申請者の自筆署名（①）が必須\n"
        "・指導医全員の自筆署名（⑦-3）が必須\n"
        "・患者イニシャル・生年月日（⑤）の記入が必須\n"
        "・本文字数（1200〜2500字）を様式内に記載すること\n"
        "・片面印刷で提出。ホチキス・クリップ・付箋等は一切使用しないこと"
    )
    set_font(footer_run, size_pt=8.5, color_rgb=(107, 114, 128))

    doc.save(output_path)
    print(f"✅ 生成完了: {output_path}  （本文字数: {char_count}字）")
    if char_count < 1200:
        print(f"   ⚠️  本文字数が1200字未満です（{char_count}字）。追記が必要です。")
    elif char_count > 2500:
        print(f"   ⚠️  本文字数が2500字を超えています（{char_count}字）。削減が必要です。")
    else:
        print(f"   ✅ 本文字数は規定範囲内です（1200〜2500字）。")


# ──────────────────────────────────────────────
# サンプルデータ（全5症例）
# ──────────────────────────────────────────────
SAMPLE_CASES = [
    {
        "case_number": 1,
        "application_date": "2025-06-01",
        "admission_type": "医療保護入院",
        "final_diagnosis": "アルコール離脱せん妄",
        "icd_code": "F10.40",
        "hospital": "○○病院 精神科",
        "patient_initial": "Y.M",
        "patient_dob": "1968-03-22",
        "admit_date": "2024-08-10",
        "discharge_date": "2024-10-05",
        "attend_start": "2024-08-10",
        "supervisor": "鈴木二郎（指定医番号：67890、指導期間：2024-08-10〜2024-10-05）",
        "restriction": "有（隔離、身体的拘束）",
        "age": 56,
        "sex": "男性",
        "chief_complaint": "「虫が見える、助けてくれ」との混乱状態で救急搬送",
        "family_history": "父：アルコール依存症。精神科的疾患の他の家族歴なし。",
        "life_history": "○○県出身。高校卒業後、建設業に従事。40歳代から飲酒量増加。妻と2人暮らし。",
        "past_history": "高血圧（降圧薬内服中）。肝機能障害（外来加療中）。精神科入院歴なし。",
        "pre_admission_course": "20代より飲酒。近年は日本酒2合/日→徐々に増加し入院前1年間はウイスキー1本/日程度。妻の勧めで断酒を試みた翌日から振戦・発汗・不眠が出現。2日目に幻視（「虫が這っている」「知らない人がいる」）・興奮状態となり、妻が救急要請。",
        "admission_situation": "搬送時は高度興奮状態、意識は変動（JCS I-2）、発熱（38.2℃）、振戦著明、発汗過多。幻視（「部屋中に虫がいる」「壁に人の顔が見える」）・錯覚あり。見当識障害（時間・場所）著明。自他への危険性高く、同意能力を欠くと判断。医療保護のために入院の必要性あり、妻の同意を得て医療保護入院とした。",
        "admission_course": "入院後、ジアゼパム静注・補液・チアミン補充を開始。興奮・幻覚は入院3日目にピークに達し、指定医として隔離・身体的拘束を命令（他患への危険・点滴抜去のおそれ）。5日目に意識清明化、7日目に隔離・拘束を解除。2週目からアルコール依存症の心理教育・AA紹介。退院後も外来通院を継続。",
        "mental_status": {
            "behavior": "興奮・多動、衣服が乱れている",
            "consciousness": "変動あり（JCS I-2）、見当識障害（時間・場所）",
            "cognition": "短期記憶障害あり",
            "mood": "恐怖・混乱・易怒性",
            "thought": "まとまりに乏しい、妄想的観念（虫に噛まれる）",
            "perception": "幻視（虫・人物）、錯覚あり",
            "risk": "自傷のおそれあり（拘束抵抗）、他患への衝動性あり",
            "insight": "病識なし、入院の必要性を認識できない",
        },
        "diagnosis_basis": "アルコール大量摂取（依存症候群：F10.2）の突然の中断2日目に出現した意識変動・幻視・自律神経過活動（発熱・振戦・発汗）は、ICD-10 F10.4（アルコール離脱状態、せん妄を伴うもの）に合致する。統合失調症・認知症・てんかんを除外。",
        "diff_diagnosis": "MRI・髄液所見異常なし（脳器質性疾患除外）。飲酒中断との時系列・臨床症状からウェルニッケ脳症との鑑別（チアミン補充で改善）。",
        "admission_basis": "妻の同意のもと医療保護入院。患者は意識変動・幻視・同意能力欠如の状態にあり、任意入院が行われる状態にないと判断。医療保護のために入院の必要性あり。",
        "restriction_detail": "入院2日目から隔離開始（他患への暴力・点滴抜去リスク。指定医命令）。3日目から身体的拘束追加（著しい不穏・自傷のおそれ切迫）。5日目に意識清明化し7日目に全行動制限解除。",
        "seisanin": "入院3日以内に退院後生活環境相談員（精神保健福祉士 佐藤）を選任。",
        "pharmacotherapy": "ジアゼパム10mg静注→経口漸減（CIWSプロトコル）。チアミン100mg/日静注3日間。ハロペリドール少量（幻覚・興奮）。",
        "nonpharmacotherapy": "アルコール依存症心理教育（全5回）・AAへの橋渡し・妻への家族教育。",
        "aftercare": "退院後週1回外来通院。断酒補助薬（アカンプロサート）開始。AA週2回参加継続。通院3ヶ月後も断酒維持。",
        "consideration": "本症例はアルコール依存症（F10.2）に離脱せん妄（F10.4）を合併した重症例。早期の身体的支援と迅速な行動制限が重篤な自傷を予防できた。退院後の地域支援（AA・外来）との連携が維持の鍵となった。",
    },
    {
        "case_number": 3,
        "application_date": "2025-06-01",
        "admission_type": "措置入院",
        "final_diagnosis": "統合失調症（妄想型）",
        "icd_code": "F20.0",
        "hospital": "△△病院 精神科",
        "patient_initial": "K.S",
        "patient_dob": "1992-11-05",
        "admit_date": "2024-05-12",
        "discharge_date": "2024-09-30",
        "attend_start": "2024-05-12",
        "supervisor": "山田三郎（指定医番号：11111、指導期間：2024-05-12〜2024-09-30）",
        "restriction": "有（隔離）",
        "age": 31,
        "sex": "男性",
        "chief_complaint": "隣人を刃物で傷つけた後、自宅に立てこもり",
        "family_history": "母：統合失調症（外来加療中）。",
        "life_history": "○○県出身。大学中退後、断続的な就労。25歳ごろから独居。社会的孤立が目立つ。",
        "past_history": "精神科受診歴あり（27歳時、外来2回で中断）。身体疾患なし。",
        "pre_admission_course": "28歳ごろから「隣人が自分の思考を盗み取っている」との確信が出現。「命令する声」が聴こえるようになり、2ヶ月前から睡眠が著しく低下。措置入院当日、隣人を包丁で切りつけ軽傷を負わせた（被害者が110番通報）。警察官が現行犯逮捕後、精神症状の疑いにより第23条通報。",
        "admission_situation": "警察官通報（精神保健福祉法第23条）により都道府県知事の命で措置診察を実施。患者は強い被害妄想（「隣人に思考を抜き取られ、意のままに操られた」）・命令する幻声（「あいつを殺せ」）を確信を持って述べ、診察に応じようとしない。意識清明、見当識保たれるが病識完全欠如。精神障害者（統合失調症）であり、かつ精神障害のために他害のおそれが著しく高いと2名の指定医が一致して判断し、措置入院とした。",
        "admission_course": "入院当日より顕著な興奮・他患への暴言。第3日目から指定医命令で個室隔離（他患への暴力行為が認められ他の方法では防ぎきれない）。リスペリドン3mg/dayから開始、漸増。3週目に妄想・幻声が和らぎ隔離解除。5週目に指定医が症状消退を判断し措置解除（医療保護入院に移行）。8週目に任意入院に形態変更後、退院支援を開始。退院後外来にて経過フォロー。",
        "mental_status": {
            "behavior": "診察室でそわそわと動き回る、警戒的",
            "consciousness": "清明",
            "cognition": "年齢相応",
            "mood": "緊張・恐怖感",
            "thought": "思考奪取・させられ体験・被害妄想（「隣人が思考を盗む・操る」）",
            "perception": "命令する声の幻聴（複数）、幻視なし",
            "risk": "他害のおそれ著しく高い（直前の傷害行為、命令幻声継続）",
            "insight": "病識なし、精神科入院の必要性を認識できない",
        },
        "diagnosis_basis": "1ヶ月以上にわたる思考奪取・させられ体験（ICD-10 F20 A基準①）および命令する第三者の声（A基準③）を認める。B基準（2症状）も満たす。経過・症状形式・家族歴から統合失調症（妄想型 F20.0）と診断。器質性疾患・薬物性精神病（尿検査陰性）を除外。",
        "diff_diagnosis": "MRI・血液検査正常（器質性除外）。覚醒剤等薬物検査陰性。双極性障害（気分エピソードを満たさない）を除外。",
        "admission_basis": "精神障害者（統合失調症 F20.0）であり、精神障害のために他害（傷害行為・命令幻声）のおそれが著しく高く、医療及び保護のために入院させなければならない状態。警察官通報（第23条）に基づく措置診察で2名の指定医が一致して自傷他害のおそれありと判定。",
        "restriction_detail": "入院3日目から個室隔離（他患への暴力行為・安全確保不可）。指定医命令。毎日診察を実施。3週目、妄想・興奮が改善し他患への危険がなくなったと判断し隔離解除。",
        "seisanin": "入院3日以内に退院後生活環境相談員（精神保健福祉士 田中）を選任。",
        "pharmacotherapy": "リスペリドン3→8mg/day（4週）。副作用（EPS・体重増加）についてインフォームドコンセントを実施。",
        "nonpharmacotherapy": "個人支持的精神療法・病識向上を目標とした認知行動的アプローチ。",
        "aftercare": "退院後外来月1回。家族（母）への心理教育。服薬継続中。症状再燃なく1年経過。",
        "consideration": "措置入院から任意入院への段階的移行を経験した症例。指定医の職権行使と人権尊重のバランス、行動制限の最小化・解除判断のタイミングが重要であった。",
    },
]


# ──────────────────────────────────────────────
# CLI
# ──────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="精神保健指定医ケースレポート .docx 生成（様式3-1 令和7年1月版準拠）"
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input",       "-i", help="入力JSONファイルのパス")
    group.add_argument("--input_dir",   "-d", help="JSONファイルが入ったディレクトリ")
    group.add_argument("--all_samples", "-s", action="store_true",
                       help="サンプルデータで docx を生成（動作確認用）")
    parser.add_argument("--output_dir", "-o", default="./output",
                        help="出力先ディレクトリ（デフォルト：./output）")
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)

    if args.all_samples:
        for case in SAMPLE_CASES:
            cn = case["case_number"]
            dx = case.get("final_diagnosis", "").replace(" ", "_").replace("（", "").replace("）", "")
            fname = f"第{cn}症例_{dx}.docx"
            out = os.path.join(args.output_dir, fname)
            generate_docx(case, out)

    elif args.input:
        with open(args.input, encoding="utf-8") as f:
            case_data = json.load(f)
        cn = case_data.get("case_number", 1)
        dx = case_data.get("final_diagnosis", "case").replace(" ", "_")
        out = os.path.join(args.output_dir, f"第{cn}症例_{dx}.docx")
        generate_docx(case_data, out)

    elif args.input_dir:
        files = [f for f in os.listdir(args.input_dir) if f.endswith(".json")]
        if not files:
            print("JSONファイルが見つかりませんでした。")
            sys.exit(1)
        for fname in sorted(files):
            path = os.path.join(args.input_dir, fname)
            with open(path, encoding="utf-8") as f:
                case_data = json.load(f)
            cn = case_data.get("case_number", 1)
            dx = case_data.get("final_diagnosis", "case").replace(" ", "_")
            out = os.path.join(args.output_dir, f"第{cn}症例_{dx}.docx")
            generate_docx(case_data, out)

    print("\n📄 生成完了。output/ ディレクトリを確認してください。")


if __name__ == "__main__":
    main()
